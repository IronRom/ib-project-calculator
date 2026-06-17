"""Generate Коммерческое предложение (КП) — Word and PDF formats matching corporate template."""
from __future__ import annotations

import datetime
import io
import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_STATIC_DIR = Path(__file__).parent.parent.parent / "static"
_SIGNATURE_IMG = _STATIC_DIR / "kp_signature.png"
_LOGO_IMG = _STATIC_DIR / "kp_logo.jpg"

_LIBERATION_DIR = Path("/usr/share/fonts/truetype/liberation")
_MAC_FONTS_DIR  = Path("/System/Library/Fonts/Supplemental")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _fmt(n: float) -> str:
    """Format number Russian-style: space thousands, comma decimal."""
    return f"{n:,.2f}".replace(",", " ").replace(".", ",")


def _rub_words(amount: float) -> str:
    """Convert ruble amount to Russian words (simplified)."""
    from math import floor

    def _hundreds(n: int) -> str:
        units = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять",
                 "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
                 "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
        tens = ["", "десять", "двадцать", "тридцать", "сорок", "пятьдесят",
                "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
        hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                    "шестьсот", "семьсот", "восемьсот", "девятьсот"]
        if n == 0:
            return "ноль"
        parts = []
        h = n // 100
        r = n % 100
        if h:
            parts.append(hundreds[h])
        if r < 20:
            if r:
                parts.append(units[r])
        else:
            parts.append(tens[r // 10])
            if r % 10:
                parts.append(units[r % 10])
        return " ".join(parts)

    total = floor(amount)
    kopecks = round((amount - total) * 100)

    billions = total // 1_000_000_000
    total %= 1_000_000_000
    millions = total // 1_000_000
    total %= 1_000_000
    thousands = total // 1_000
    remainder = total % 1_000

    parts: list[str] = []

    if billions:
        w = _hundreds(billions)
        parts.append(f"{w} {'миллиард' if billions == 1 else 'миллиардов'}")

    if millions:
        w = _hundreds(millions)
        suffix = ("миллион" if millions % 10 == 1 and millions % 100 != 11 else
                  "миллиона" if 2 <= millions % 10 <= 4 and not (11 <= millions % 100 <= 14) else
                  "миллионов")
        parts.append(f"{w} {suffix}")

    if thousands:
        units_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять",
                   "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
                   "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
        tens_f = ["", "десять", "двадцать", "тридцать", "сорок", "пятьдесят",
                  "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
        hundreds_f = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                      "шестьсот", "семьсот", "восемьсот", "девятьсот"]
        th_parts = []
        h = thousands // 100
        r = thousands % 100
        if h:
            th_parts.append(hundreds_f[h])
        if r < 20:
            if r:
                th_parts.append(units_f[r])
        else:
            th_parts.append(tens_f[r // 10])
            if r % 10:
                th_parts.append(units_f[r % 10])
        w = " ".join(th_parts) if th_parts else ""
        suffix = ("тысяча" if thousands % 10 == 1 and thousands % 100 != 11 else
                  "тысячи" if 2 <= thousands % 10 <= 4 and not (11 <= thousands % 100 <= 14) else
                  "тысяч")
        parts.append(f"{w} {suffix}".strip())

    if remainder:
        parts.append(_hundreds(remainder))

    words = " ".join(parts) if parts else "ноль"
    words = words[0].upper() + words[1:] if words else "Ноль"
    return f"{words} рублей {kopecks:02d} копеек"


def _set_cell_borders(cell, color: str = "000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'left', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_run(para: Any, text: str, bold: bool = False, italic: bool = False,
             underline: bool = False, size: int = 12) -> Any:
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return run


# ── Word generator ────────────────────────────────────────────────────────────

def generate_kp_word(
    project_name: str,
    stage: str,
    result: dict[str, Any],
    tz_object_name: str = "",
    company_name: str = "",
) -> bytes:
    """Generate КП Word document matching the corporate template."""
    obj_name = tz_object_name.strip() or project_name

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(1.5)

    # ── Logo in header ────────────────────────────────────────────────────────
    if _LOGO_IMG.exists():
        header = doc.sections[0].header
        hdr_para = header.paragraphs[0]
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hdr_para.add_run()
        run.add_picture(str(_LOGO_IMG), width=Cm(2.5))

    vat_rate       = (result.get("vat_rate") or 22) / 100
    total_with_vat = float(result.get("total_with_vat") or 0)
    cost_with_stage = float(result.get("cost_with_stage") or 0)
    vat_amount     = float(result.get("vat_amount") or 0)
    vat_pct        = int(vat_rate * 100)

    today = datetime.date.today()
    date_str = f"{today.day:02d}.{today.month:02d}.{today.year} г."

    # ── Исх. line ─────────────────────────────────────────────────────────────
    ish_p = doc.add_paragraph()
    ish_p.paragraph_format.space_after = Pt(6)
    _add_run(ish_p, f"Исх.: КП от {date_str}", size=12)

    # ── Касательно ────────────────────────────────────────────────────────────
    kas_p = doc.add_paragraph()
    kas_p.paragraph_format.space_after = Pt(0)
    _add_run(kas_p, "Касательно:", italic=True, size=12)

    # Italic body of касательно
    body_ref_p = doc.add_paragraph()
    body_ref_p.paragraph_format.space_after = Pt(6)
    _add_run(body_ref_p,
             f"Коммерческое предложение по объекту: «{obj_name}»",
             italic=True, size=12)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    _add_run(title_p, "Коммерческое предложение", bold=True, size=14)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Body paragraph ────────────────────────────────────────────────────────
    stage_labels = {
        "П":   "проектной документации стадии ПД",
        "Р":   "рабочей документации стадии РД",
        "П+Р": "проектной документации стадий ПД и РД",
    }
    stage_label = stage_labels.get(stage, "проектной и рабочей документации")
    total_words = _rub_words(total_with_vat)

    body_p = doc.add_paragraph()
    body_p.paragraph_format.space_after = Pt(6)
    body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_run(body_p,
             f"      Компания ООО «Интеллект-Строй» направляет Вам коммерческое "
             f"предложение на разработку {stage_label} объекта: ")
    _add_run(body_p, f"«{obj_name}»", bold=True)
    _add_run(body_p, " стоимостью ")
    _add_run(body_p, f"{_fmt(total_with_vat)} руб.", bold=True)
    _add_run(body_p,
             f" ({total_words}), в т.ч. НДС ({vat_pct}%) {_fmt(vat_amount)} руб.")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Cost table ────────────────────────────────────────────────────────────
    lines: list[tuple[str, float]] = []
    if stage == "П+Р":
        pd_cost = cost_with_stage * 0.4 * (1 + vat_rate)
        rd_cost = cost_with_stage * 0.6 * (1 + vat_rate)
        lines = [
            ("Разработка проектной документации", pd_cost),
            ("Разработка рабочей документации",   rd_cost),
        ]
    elif stage == "П":
        lines = [("Разработка проектной документации", total_with_vat)]
    else:
        lines = [("Разработка рабочей документации", total_with_vat)]

    table = doc.add_table(rows=0, cols=3)
    col_widths = [Cm(1.5), Cm(10.5), Cm(4.5)]

    # Header row
    hdr = table.add_row()
    for ci, (cell, text, w) in enumerate(zip(hdr.cells, [
        "№\nп/п", "Наименование", f"Сумма, с НДС\n{vat_pct}%"
    ], col_widths)):
        cell.width = w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        _set_cell_borders(cell)
        _set_cell_shading(cell, "D9D9D9")

    # Data rows
    for num, (name, cost) in enumerate(lines, 1):
        row = table.add_row()
        for ci, (cell, text) in enumerate(zip(row.cells, [str(num), name, _fmt(cost)])):
            cell.width = col_widths[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
            _set_cell_borders(cell)

    # НДС row
    vat_row = table.add_row()
    for ci, (cell, text) in enumerate(zip(vat_row.cells, ["", f"НДС {vat_pct}%", _fmt(vat_amount)])):
        cell.width = col_widths[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci == 1 else (WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(text)
        r.italic = (ci == 1)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        _set_cell_borders(cell)

    # Total row
    tot_row = table.add_row()
    for ci, (cell, text) in enumerate(zip(tot_row.cells, ["", f"Итого с НДС {vat_pct}%", _fmt(total_with_vat)])):
        cell.width = col_widths[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci == 1 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        _set_cell_borders(cell)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Validity bullet ───────────────────────────────────────────────────────
    bullet = doc.add_paragraph(style='List Bullet')
    bullet.paragraph_format.space_after = Pt(12)
    r = bullet.add_run(
        "Данное коммерческое предложение действует в течении 15 (пятнадцати) рабочих дней."
    )
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    # ── Signature block ───────────────────────────────────────────────────────
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.style = "Table Grid"

    for cell in sig_table.rows[0].cells:
        _remove_cell_borders(cell)

    left_cell = sig_table.rows[0].cells[0]
    left_cell.width = Cm(5)
    for line in ["С уважением,", "Генеральный директор", "ООО «Интеллект-строй»"]:
        p = left_cell.add_paragraph()
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    mid_cell = sig_table.rows[0].cells[1]
    mid_cell.width = Cm(8)
    if _SIGNATURE_IMG.exists():
        img_para = mid_cell.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(_SIGNATURE_IMG), width=Cm(7.5))

    right_cell = sig_table.rows[0].cells[2]
    right_cell.width = Cm(4)
    name_para = right_cell.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    name_para.paragraph_format.space_before = Pt(28)
    r = name_para.add_run("И. А. Подопригора")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PDF generator ─────────────────────────────────────────────────────────────

_FONTS_REGISTERED: dict[str, str] = {}


def _register_fonts() -> str:
    """Register a Cyrillic-capable serif font; return family name."""
    if "family" in _FONTS_REGISTERED:
        return _FONTS_REGISTERED["family"]

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    candidates = [
        # Linux / Docker (fonts-liberation)
        (
            _LIBERATION_DIR / "LiberationSerif-Regular.ttf",
            _LIBERATION_DIR / "LiberationSerif-Bold.ttf",
            _LIBERATION_DIR / "LiberationSerif-Italic.ttf",
            _LIBERATION_DIR / "LiberationSerif-BoldItalic.ttf",
        ),
        # macOS
        (
            _MAC_FONTS_DIR / "Times New Roman.ttf",
            _MAC_FONTS_DIR / "Times New Roman Bold.ttf",
            _MAC_FONTS_DIR / "Times New Roman Italic.ttf",
            _MAC_FONTS_DIR / "Times New Roman Bold Italic.ttf",
        ),
    ]

    for regular, bold, italic, bold_italic in candidates:
        if regular.exists() and bold.exists():
            try:
                pdfmetrics.registerFont(TTFont("TNR",           str(regular)))
                pdfmetrics.registerFont(TTFont("TNR-Bold",      str(bold)))
                pdfmetrics.registerFont(TTFont("TNR-Italic",    str(italic)))
                pdfmetrics.registerFont(TTFont("TNR-BoldItalic", str(bold_italic)))
                registerFontFamily("TNR", normal="TNR", bold="TNR-Bold",
                                   italic="TNR-Italic", boldItalic="TNR-BoldItalic")
                _FONTS_REGISTERED["family"] = "TNR"
                return "TNR"
            except Exception:
                continue

    _FONTS_REGISTERED["family"] = "Times-Roman"
    return "Times-Roman"


def generate_kp_pdf(
    project_name: str,
    stage: str,
    result: dict[str, Any],
    tz_object_name: str = "",
    company_name: str = "",
) -> bytes:
    """Generate КП as PDF matching the corporate template."""
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, Image as RLImage
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib import colors

    font = _register_fonts()
    bold_font  = f"{font}-Bold"   if font != "Times-Roman" else "Times-Bold"
    ital_font  = f"{font}-Italic" if font != "Times-Roman" else "Times-Italic"

    obj_name       = tz_object_name.strip() or project_name
    vat_rate       = (result.get("vat_rate") or 22) / 100
    total_with_vat = float(result.get("total_with_vat") or 0)
    cost_with_stage = float(result.get("cost_with_stage") or 0)
    vat_amount     = float(result.get("vat_amount") or 0)
    vat_pct        = int(vat_rate * 100)

    today    = datetime.date.today()
    date_str = f"{today.day:02d}.{today.month:02d}.{today.year} г."

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3*cm, rightMargin=1.5*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )
    W = doc.width

    def style(align=TA_LEFT, size=12, bold=False, italic=False, first=0, leading=17, color=colors.black):
        fn = (bold_font if bold else ital_font if italic else font)
        return ParagraphStyle("_", fontName=fn, fontSize=size, leading=leading,
                              alignment=align, firstLineIndent=first,
                              spaceAfter=0, textColor=color)

    story = []

    # ── Header row: Исх. | Logo ───────────────────────────────────────────────
    isх_text = Paragraph(f"Исх.: КП от {date_str}", style())
    logo_cell: Any = ""
    if _LOGO_IMG.exists():
        logo_cell = RLImage(str(_LOGO_IMG), width=2.5*cm, height=2.5*cm)

    hdr_data = [[isх_text, logo_cell]]
    hdr_table = Table(hdr_data, colWidths=[W - 3*cm, 3*cm])
    hdr_table.setStyle(TableStyle([
        ("ALIGN",   (1, 0), (1, 0), "RIGHT"),
        ("VALIGN",  (0, 0), (-1, -1), "TOP"),
        ("LINEBELOW", (0, 0), (-1, -1), 0, colors.white),
        ("BOX",     (0, 0), (-1, -1), 0, colors.white),
        ("INNERGRID", (0, 0), (-1, -1), 0, colors.white),
    ]))
    story.append(hdr_table)
    story.append(Spacer(1, 0.5*cm))

    # ── Касательно ────────────────────────────────────────────────────────────
    story.append(Paragraph("Касательно:", style(italic=True)))
    story.append(Paragraph(
        f"Коммерческое предложение по объекту: «{obj_name}»",
        style(italic=True, leading=18),
    ))
    story.append(Spacer(1, 0.8*cm))

    # ── Title ─────────────────────────────────────────────────────────────────
    story.append(Paragraph("Коммерческое предложение", style(align=TA_CENTER, size=14, bold=True, leading=20)))
    story.append(Spacer(1, 0.5*cm))

    # ── Body paragraph ────────────────────────────────────────────────────────
    stage_labels = {
        "П":   "проектной документации стадии ПД",
        "Р":   "рабочей документации стадии РД",
        "П+Р": "проектной документации стадий ПД и РД",
    }
    stage_label = stage_labels.get(stage, "проектной и рабочей документации")
    total_words = _rub_words(total_with_vat)

    body_html = (
        f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
        f"Компания ООО «Интеллект-Строй» направляет Вам коммерческое предложение "
        f"на разработку {stage_label} объекта: "
        f"<b>«{obj_name}»</b> стоимостью "
        f"<b>{_fmt(total_with_vat)} руб.</b> ({total_words}), "
        f"в т.ч. НДС ({vat_pct}%) {_fmt(vat_amount)} руб."
    )
    story.append(Paragraph(body_html, style(align=TA_JUSTIFY, leading=18)))
    story.append(Spacer(1, 0.5*cm))

    # ── Cost table ────────────────────────────────────────────────────────────
    lines: list[tuple[str, float]] = []
    if stage == "П+Р":
        pd_cost = cost_with_stage * 0.4 * (1 + vat_rate)
        rd_cost = cost_with_stage * 0.6 * (1 + vat_rate)
        lines = [
            ("Разработка проектной документации", pd_cost),
            ("Разработка рабочей документации",   rd_cost),
        ]
    elif stage == "П":
        lines = [("Разработка проектной документации", total_with_vat)]
    else:
        lines = [("Разработка рабочей документации", total_with_vat)]

    col_w = [1.5*cm, 10.5*cm, 4.5*cm]
    th_style = style(align=TA_CENTER, size=11, bold=True)
    td_style = style(size=11)
    td_right = style(align=TA_RIGHT, size=11)

    table_data = [
        [
            Paragraph("№\nп/п", th_style),
            Paragraph("Наименование", th_style),
            Paragraph(f"Сумма, с НДС\n{vat_pct}%", th_style),
        ]
    ]
    for num, (name, cost) in enumerate(lines, 1):
        table_data.append([
            Paragraph(str(num), style(align=TA_CENTER, size=11)),
            Paragraph(name, td_style),
            Paragraph(_fmt(cost), td_right),
        ])
    table_data.append([
        Paragraph("", td_style),
        Paragraph(f"НДС {vat_pct}%", style(align=TA_RIGHT, size=11, italic=True)),
        Paragraph(_fmt(vat_amount), td_right),
    ])
    table_data.append([
        Paragraph("", td_style),
        Paragraph(f"Итого с НДС {vat_pct}%", style(align=TA_RIGHT, size=11, bold=True)),
        Paragraph(_fmt(total_with_vat), style(align=TA_RIGHT, size=11, bold=True)),
    ])

    cost_table = Table(table_data, colWidths=col_w)
    cost_table.setStyle(TableStyle([
        ("BOX",       (0, 0), (-1, -1), 0.5, colors.black),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9D9D9")),
        ("VALIGN",    (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",  (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(cost_table)
    story.append(Spacer(1, 0.4*cm))

    # ── Validity bullet ───────────────────────────────────────────────────────
    story.append(Paragraph(
        "• Данное коммерческое предложение действует в течении 15 (пятнадцати) рабочих дней.",
        style(leading=18),
    ))
    story.append(Spacer(1, 1.0*cm))

    # ── Signature block ───────────────────────────────────────────────────────
    left_text = Paragraph(
        "С уважением,<br/>Генеральный директор<br/>ООО «Интеллект-строй»",
        style(bold=True, leading=18),
    )
    sig_img: Any = ""
    if _SIGNATURE_IMG.exists():
        sig_img = RLImage(str(_SIGNATURE_IMG), width=6*cm, height=4*cm)

    right_text = Paragraph("И. А. Подопригора", style(bold=True, align=TA_RIGHT))

    sig_data = [[left_text, sig_img, right_text]]
    # 5.5 + 6 + 5 = 16.5cm = A4 - margins
    sig_table = Table(sig_data, colWidths=[5.5*cm, 6*cm, 5*cm])
    sig_table.setStyle(TableStyle([
        ("BOX",       (0, 0), (-1, -1), 0, colors.white),
        ("INNERGRID", (0, 0), (-1, -1), 0, colors.white),
        ("VALIGN",    (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN",     (2, 0), (2, 0), "RIGHT"),
    ]))
    story.append(sig_table)

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()
